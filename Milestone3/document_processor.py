# Document processing functionality for the AI-Based Smart File Assistant
import time
from config import Config
from pinecone_manager import get_user_pinecone_index, create_user_pinecone_index
from database import get_db
from embeddings import generate_embedding_miniLM

# Try to import LangChain components, fall back to basic processing if not available
try:
    from langchain_community.document_loaders import (
        PyPDFLoader,
        Docx2txtLoader,
        UnstructuredHTMLLoader,
        TextLoader
    )
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
    print("✅ LangChain components loaded successfully")
except ImportError as e:
    print(f"⚠️  LangChain not fully available: {e}")
    LANGCHAIN_AVAILABLE = False

def extract_text_from_file_langchain(file_path, filename):
    """Extract text using LangChain document loaders with fallback to basic extraction"""
    try:
        file_ext = filename.lower().split('.')[-1]
        
        if LANGCHAIN_AVAILABLE:
            # Try LangChain loaders first
            try:
                if file_ext == 'pdf':
                    loader = PyPDFLoader(file_path)
                elif file_ext in ['doc', 'docx']:
                    loader = Docx2txtLoader(file_path)
                elif file_ext == 'html':
                    loader = UnstructuredHTMLLoader(file_path)
                elif file_ext == 'txt':
                    loader = TextLoader(file_path, encoding='utf-8')
                else:
                    print(f"❌ Unsupported file type: {file_ext}")
                    return None
                
                # Load documents
                documents = loader.load()
                print(f"✅ Loaded {len(documents)} document(s) from {filename} using LangChain")
                return documents
                
            except Exception as e:
                print(f"⚠️  LangChain loader failed for {filename}: {e}")
                # Fall through to basic extraction
        
        # Fallback to basic text extraction
        print(f"🔄 Using basic text extraction for {filename}")
        text_content = ""
        
        if file_ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
        elif file_ext == 'pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
            except ImportError:
                print("⚠️  PyPDF2 not installed. Install with: pip install PyPDF2")
                return None
        elif file_ext in ['doc', 'docx']:
            try:
                from docx import Document
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
            except ImportError:
                print("⚠️  python-docx not installed. Install with: pip install python-docx")
                return None
        else:
            print(f"❌ Unsupported file type: {file_ext}")
            return None
        
        # Create a document-like object for consistency
        class SimpleDocument:
            def __init__(self, content, source):
                self.page_content = content
                self.metadata = {'source': source}
        
        return [SimpleDocument(text_content.strip(), filename)]
        
    except Exception as e:
        print(f"❌ Error loading {filename}: {str(e)}")
        return None

def split_documents_langchain(documents):
    """Split documents using RecursiveCharacterTextSplitter with fallback"""
    try:
        if LANGCHAIN_AVAILABLE:
            # Try LangChain splitter
            try:
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=Config.CHUNK_CONFIG["chunk_size"],
                    chunk_overlap=Config.CHUNK_CONFIG["chunk_overlap"]
                )
                
                chunks = splitter.split_documents(documents)
                print(f"✅ Split into {len(chunks)} chunks using LangChain")
                return chunks
            except Exception as e:
                print(f"⚠️  LangChain splitter failed: {e}")
        
        # Fallback to basic splitting
        print(f"🔄 Using basic text splitting")
        chunks = []
        
        for doc in documents:
            text = doc.page_content
            chunk_size = Config.CHUNK_CONFIG["chunk_size"]
            overlap = Config.CHUNK_CONFIG["chunk_overlap"]
            
            start = 0
            chunk_index = 0
            
            while start < len(text):
                end = start + chunk_size
                
                # Try to break at sentence or word boundary
                if end < len(text):
                    sentence_end = text.rfind('.', start, end)
                    if sentence_end > start:
                        end = sentence_end + 1
                    else:
                        word_end = text.rfind(' ', start, end)
                        if word_end > start:
                            end = word_end
                
                chunk_text = text[start:end].strip()
                if chunk_text and len(chunk_text) >= Config.CHUNK_CONFIG["min_chunk_length"]:
                    # Create chunk object
                    class SimpleChunk:
                        def __init__(self, content, metadata):
                            self.page_content = content
                            self.metadata = metadata
                    
                    chunk_metadata = doc.metadata.copy()
                    chunk_metadata['chunk_index'] = chunk_index
                    chunks.append(SimpleChunk(chunk_text, chunk_metadata))
                    chunk_index += 1
                
                start = end - overlap
                if start >= len(text):
                    break
        
        print(f"✅ Split into {len(chunks)} chunks using basic splitting")
        return chunks
        
    except Exception as e:
        print(f"❌ Error splitting documents: {str(e)}")
        return []

def process_document_for_user(user_id, file_path, filename):
    """Process a document using the same pipeline as Milestone2 with enhanced record tracking"""
    print(f"🔧 DEBUG: process_document_for_user called with:")
    print(f"   user_id: {user_id}")
    print(f"   file_path: {file_path}")
    print(f"   filename: {filename}")
    
    try:
        print(f"\n🚀 Starting document processing for user {user_id}")
        print(f"📄 Processing document: {filename}")
        
        # Get user's individual Pinecone index
        index = get_user_pinecone_index(user_id)
        if not index:
            print(f"❌ No Pinecone index found for user {user_id}, attempting to create one...")
            
            # Get user email for index creation
            conn = get_db()
            cursor = conn.cursor()
            cursor.execute("SELECT email FROM users WHERE id = ?", (user_id,))
            user_result = cursor.fetchone()
            conn.close()
            
            if user_result:
                email = user_result[0]
                index_name = create_user_pinecone_index(user_id, email)
                
                if index_name:
                    # Update user record with new index name
                    conn = get_db()
                    cursor = conn.cursor()
                    cursor.execute("UPDATE users SET index_name = ? WHERE id = ?", (index_name, user_id))
                    conn.commit()
                    conn.close()
                    
                    # Try to get the index again
                    index = get_user_pinecone_index(user_id)
                    if index:
                        print(f"✅ Created and connected to new Pinecone index for user {user_id}")
                    else:
                        print(f"❌ Failed to connect to newly created index for user {user_id}")
                        return {"success": False, "error": "Failed to connect to Pinecone index"}
                else:
                    print(f"❌ Failed to create Pinecone index for user {user_id}")
                    return {"success": False, "error": "Failed to create Pinecone index"}
            else:
                print(f"❌ Could not find user {user_id} in database")
                return {"success": False, "error": "User not found"}
        else:
            print(f"✅ Connected to existing Pinecone index for user {user_id}")
        
        # Step 1: Load documents using LangChain (same as Milestone2)
        documents = extract_text_from_file_langchain(file_path, filename)
        if not documents:
            print(f"⚠️  Could not load document {filename}")
            return {"success": False, "error": "Could not load document"}
        
        print(f"📝 Loaded {len(documents)} document(s) from {filename}")
        
        # Step 2: Split documents into chunks (same as Milestone2)
        chunks = split_documents_langchain(documents)
        if not chunks:
            print(f"⚠️  Could not split document {filename}")
            return {"success": False, "error": "Could not split document"}
        
        print(f"🔪 Split document into {len(chunks)} chunks")
        
        # Apply safety limit
        original_chunk_count = len(chunks)
        if len(chunks) > Config.CHUNK_CONFIG["max_chunks_per_doc"]:
            print(f"⚠️  Too many chunks ({len(chunks)}), limiting to {Config.CHUNK_CONFIG['max_chunks_per_doc']}")
            chunks = chunks[:Config.CHUNK_CONFIG["max_chunks_per_doc"]]
        
        total_chunks = len(chunks)
        print(f"📊 Total chunks to process: {total_chunks}")
        
        # Step 3: Generate embeddings and prepare vectors (same as Milestone2)
        vectors_to_upsert = []
        successful_chunks = 0
        failed_chunks = 0
        
        for i, chunk in enumerate(chunks):
            print(f"🔄 Processing chunk {i+1}/{total_chunks}")
            
            # Generate embedding using MiniLM (same as Milestone2)
            embedding = generate_embedding_miniLM(chunk.page_content)
            if embedding is None:
                print(f"⚠️  Failed to generate embedding for chunk {i+1}")
                failed_chunks += 1
                continue
            
            # Create unique ID for this chunk
            timestamp = int(time.time())
            chunk_id = f"user_{user_id}_doc_{timestamp}_{i}"
            
            # Prepare metadata (enhanced with more info)
            metadata = {
                "user_id": user_id,
                "filename": filename,
                "chunk_index": i,
                "total_chunks": total_chunks,
                "text": chunk.page_content[:1000],  # Store first 1000 chars for preview
                "upload_time": timestamp,
                "file_path": file_path,
                "chunk_size": len(chunk.page_content),
                "source": chunk.metadata.get('source', filename),
                "page": chunk.metadata.get('page', 0)
            }
            
            vectors_to_upsert.append((chunk_id, embedding, metadata))
            successful_chunks += 1
            print(f"✅ Prepared chunk {i+1} for upload (ID: {chunk_id})")
        
        print(f"📊 Successfully prepared {successful_chunks}/{total_chunks} chunks")
        print(f"📊 Failed chunks: {failed_chunks}")
        
        # Step 4: Upload to Pinecone in batches (same as Milestone2)
        uploaded_vectors = 0
        if vectors_to_upsert:
            batch_size = 100
            total_batches = (len(vectors_to_upsert) - 1) // batch_size + 1
            
            for i in range(0, len(vectors_to_upsert), batch_size):
                batch = vectors_to_upsert[i:i + batch_size]
                batch_num = i // batch_size + 1
                
                print(f"📤 Uploading batch {batch_num}/{total_batches} ({len(batch)} vectors)")
                
                try:
                    index.upsert(vectors=batch)
                    uploaded_vectors += len(batch)
                    print(f"✅ Successfully uploaded batch {batch_num}")
                except Exception as e:
                    print(f"❌ Failed to upload batch {batch_num}: {str(e)}")
                    return {"success": False, "error": f"Failed to upload batch {batch_num}: {str(e)}"}
        
        # Final statistics
        processing_stats = {
            "success": True,
            "filename": filename,
            "original_documents": len(documents),
            "original_chunks": original_chunk_count,
            "processed_chunks": total_chunks,
            "successful_chunks": successful_chunks,
            "failed_chunks": failed_chunks,
            "uploaded_vectors": uploaded_vectors,
            "pinecone_records": uploaded_vectors,
            "chunk_size_config": Config.CHUNK_CONFIG["chunk_size"],
            "chunk_overlap_config": Config.CHUNK_CONFIG["chunk_overlap"]
        }
        
        print(f"🎉 Successfully processed {filename}!")
        print(f"📊 Final stats:")
        print(f"   - Original documents: {len(documents)}")
        print(f"   - Original chunks: {original_chunk_count}")
        print(f"   - Processed chunks: {total_chunks}")
        print(f"   - Successful chunks: {successful_chunks}")
        print(f"   - Failed chunks: {failed_chunks}")
        print(f"   - Uploaded vectors to Pinecone: {uploaded_vectors}")
        
        return processing_stats
        
    except Exception as e:
        print(f"❌ Error processing document {filename} for user {user_id}: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}